"""
Resume Reactor - Export Service
Generates ATS-friendly DOCX and PDF outputs
"""
import os
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

from config import TEMP_DIR


async def export_resume(
    resume_data: Dict[str, Any],
    format: str
) -> str:
    """
    Export resume to specified format (docx or pdf)
    """
    resume_id = resume_data.get("file_path", "").split("\\")[-1].split(".")[0]
    output_filename = f"{resume_id}_optimized.{format}"
    output_path = os.path.join(TEMP_DIR, output_filename)
    
    parsed = resume_data.get("parsed", {})
    sections = parsed.get("sections", {})
    
    if format == "docx":
        create_docx(output_path, sections, parsed.get("text", ""))
    else:
        create_pdf(output_path, sections, parsed.get("text", ""))
    
    return output_path


def create_docx(output_path: str, sections: Dict[str, str], full_text: str):
    """
    Create ATS-friendly DOCX resume
    Simple formatting, no tables or graphics for maximum ATS compatibility
    """
    doc = Document()
    
    # Set narrow margins for more content space
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add contact section (usually at top)
    if sections.get("contact"):
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in sections["contact"].split('\n')[:4]:
            run = contact.add_run(line + '\n')
            run.font.size = Pt(10)
    
    # Section order for ATS
    section_order = ["summary", "experience", "education", "skills", "projects", "certifications"]
    section_titles = {
        "summary": "PROFESSIONAL SUMMARY",
        "experience": "PROFESSIONAL EXPERIENCE",
        "education": "EDUCATION",
        "skills": "SKILLS",
        "projects": "PROJECTS",
        "certifications": "CERTIFICATIONS"
    }
    
    for section_key in section_order:
        content = sections.get(section_key, "")
        if content:
            # Add section header
            header = doc.add_paragraph()
            header_run = header.add_run(section_titles.get(section_key, section_key.upper()))
            header_run.bold = True
            header_run.font.size = Pt(12)
            
            # Add section content
            for line in content.split('\n'):
                if line.strip():
                    para = doc.add_paragraph()
                    para.paragraph_format.space_after = Pt(2)
                    run = para.add_run(line)
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
    
    doc.save(output_path)


def create_pdf(output_path: str, sections: Dict[str, str], full_text: str):
    """
    Create clean, readable PDF resume
    """
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    header_style = ParagraphStyle(
        'Header',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=6,
        spaceBefore=12,
        textColor='#1a1a1a'
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=3,
        leading=14
    )
    
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=10,
        alignment=1,  # Center
        spaceAfter=12
    )
    
    story = []
    
    # Contact info
    if sections.get("contact"):
        contact_text = sections["contact"].replace('\n', '<br/>')
        story.append(Paragraph(contact_text, contact_style))
        story.append(Spacer(1, 0.2*inch))
    
    # Sections
    section_order = ["summary", "experience", "education", "skills", "projects", "certifications"]
    section_titles = {
        "summary": "PROFESSIONAL SUMMARY",
        "experience": "PROFESSIONAL EXPERIENCE",
        "education": "EDUCATION",
        "skills": "SKILLS",
        "projects": "PROJECTS",
        "certifications": "CERTIFICATIONS"
    }
    
    for section_key in section_order:
        content = sections.get(section_key, "")
        if content:
            # Section header
            story.append(Paragraph(section_titles.get(section_key, section_key.upper()), header_style))
            
            # Section content
            for line in content.split('\n'):
                if line.strip():
                    # Escape special characters for reportlab
                    safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(safe_line, body_style))
    
    doc.build(story)


def get_ats_safe_template() -> Dict[str, Any]:
    """
    Return template structure for ATS-safe resume
    """
    return {
        "fonts": ["Calibri", "Arial", "Times New Roman"],
        "recommended_sections": [
            "Contact Information",
            "Professional Summary",
            "Work Experience",
            "Education",
            "Skills"
        ],
        "avoid": [
            "Tables",
            "Text boxes",
            "Headers/Footers",
            "Images/graphics",
            "Columns",
            "Fancy fonts"
        ],
        "tips": [
            "Use standard section headings",
            "List experience in reverse chronological order",
            "Include keywords from the job description",
            "Use bullet points for achievements",
            "Quantify results where possible"
        ]
    }
