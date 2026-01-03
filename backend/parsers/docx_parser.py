"""
Resume Reactor - DOCX Parser
Extracts text and images from DOCX resumes
"""
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import re
from typing import Dict, Any


def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    Parse a DOCX resume and extract text, sections, and image count.
    """
    text_content = ""
    images_count = 0
    
    try:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + "\n"
        
        # Count images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images_count += 1
                
    except Exception as e:
        print(f"DOCX parsing error: {e}")
    
    # Parse sections from text
    sections = extract_sections(text_content)
    
    return {
        "text": text_content.strip(),
        "sections": sections,
        "images_count": images_count
    }


def extract_sections(text: str) -> Dict[str, str]:
    """
    Extract resume sections based on common headings
    """
    sections = {
        "contact": "",
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
        "certifications": "",
        "other": ""
    }
    
    # Common section headers (case insensitive)
    section_patterns = {
        "summary": r"(?:professional\s+)?summary|objective|profile|about\s*me",
        "experience": r"(?:work\s+)?experience|employment|work\s*history|professional\s*experience",
        "education": r"education|academic|qualifications|degrees?",
        "skills": r"skills?|technical\s*skills?|competenc(?:y|ies)|technologies|expertise",
        "projects": r"projects?|portfolio|work\s*samples?",
        "certifications": r"certifications?|certificates?|licenses?|credentials?"
    }
    
    lines = text.split('\n')
    current_section = "contact"
    section_content = []
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        # Check if this line is a section header
        matched_section = None
        for section_name, pattern in section_patterns.items():
            if re.match(f"^{pattern}.*$", line_stripped, re.IGNORECASE):
                matched_section = section_name
                break
        
        if matched_section:
            if section_content:
                sections[current_section] = '\n'.join(section_content)
            current_section = matched_section
            section_content = []
        else:
            section_content.append(line_stripped)
    
    if section_content:
        sections[current_section] = '\n'.join(section_content)
    
    return sections


def extract_formatting_info(file_path: str) -> Dict[str, Any]:
    """
    Extract formatting information for ATS compatibility analysis
    """
    formatting = {
        "has_tables": False,
        "has_headers": False,
        "has_footers": False,
        "font_count": 0,
        "has_columns": False
    }
    
    try:
        doc = Document(file_path)
        
        formatting["has_tables"] = len(doc.tables) > 0
        
        # Check for headers/footers
        for section in doc.sections:
            if section.header.paragraphs:
                formatting["has_headers"] = True
            if section.footer.paragraphs:
                formatting["has_footers"] = True
        
        # Count unique fonts (simplified)
        fonts = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts.add(run.font.name)
        formatting["font_count"] = len(fonts)
        
    except Exception as e:
        print(f"Formatting analysis error: {e}")
    
    return formatting
